import re
import time
from appium import webdriver
import pandas as pd
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import random
from collections import deque
import json
from appium.webdriver.common.touch_action import TouchAction
from adb_shell.adb_device import AdbDeviceTcp, AdbDeviceUsb
from selenium.common.exceptions import NoSuchElementException
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from time import sleep
from config import *



class Moments():
    def __init__(self):
        """
        初始化
        """
        # 驱动配置
        self.desired_caps = {
            'platformName': PLATFORM,
            'deviceName': DEVICE_NAME,
            'appPackage': APP_PACKAGE,
            'appActivity': APP_ACTIVITY,
            'noReset': NORESET
        }
        self.driver = webdriver.Remote(DRIVER_SERVER, self.desired_caps)
        self.wait = WebDriverWait(self.driver, TIMEOUT)
        self.adb_device = AdbDeviceTcp('127.0.0.1', 5555, default_transport_timeout_s=9.)
        self.adb_device.connect()
        self.now_year = time.strftime("%Y年", time.localtime())
        self.start_time = time.mktime(time.strptime(STARTTIME, "%Y年%m月%d日"))
        print(self.start_time)

    def contact(self):
        """
        定位通讯录
        :return:
        """
#        tab = self.wait.until(EC.presence_of_element_located((By.XPATH, '//android.widget.ImageView[@resource-id="com.tencent.mm:id/f2s"][3]')))
        tab = self.wait.until(EC.presence_of_all_elements_located((By.XPATH, '//*[@resource-id="com.tencent.mm:id/f2a"]')))[1]
        tab.click()
        # 找到目标
#        sleep(SCROLL_SLEEP_TIME)
        while True:
            try:
                selectf = self.driver.find_element(By.XPATH, '//*[@text="' + FRIEND + '"]')
                selectf.click()
                break
            except NoSuchElementException:
                self.driver.swipe(FLICK_START_X, FLICK_START_Y + 1000, FLICK_START_X, FLICK_START_Y, 1000)
                pass
#        selectf = self.wait.until(EC.presence_of_element_located((By.XPATH, '//*[@text="' + FRIEND + '"]')))
#        selectf.click()
        # 进入朋友圈
        # f = self.wait.until(EC.presence_of_element_located((By.XPATH, '//*[@resource-id="com.tencent.mm:id/iwg"][3]')))
        # f = self.wait.until(EC.presence_of_element_located((By.ID, 'com.tencent.mm:id/iwg')))
        f = self.wait.until(EC.presence_of_all_elements_located((By.ID, 'com.tencent.mm:id/iwg')))[2]
        sleep(0.5)
        f.click()

    def w_docx_content(self, document, art_list):
        """
        写入段落
        :param document:
        :param art_list:
        :return:
        """
        for p in art_list:
            pg = document.add_paragraph()
            # 设置内容
            pg.text = p.replace(" ", "")
            # 设置字号
            pg.style.font.size = Pt(14)
            pg.paragraph_format.first_line_indent = pg.style.font.size * 2
            # 设置英文字符字体
            pg.style.font.name = '黑体'
            # 设置中文字符字体
            pg.style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    def w_docx_title(self, document, art_title):
        """
        写入标题
        :param document:
        :param art_title:
        :return:
        """
        run = document.add_heading("", level=0).add_run(art_title.lstrip())
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(14)

    def tranform_time(self, art_time):
        """
        修改时间格式
        :param art_time: 传入日期
        :return: 日期和年份
        """
        if art_time == "今天":
            art_r_time = time.strftime("%m月%d日", time.localtime())
        elif art_time == "昨天":
            art_r_time = time.strftime("%m月%d日", time.localtime(time.mktime(time.localtime()) - (3600 * 24)))
        elif art_time == "前天":
            art_r_time = time.strftime("%m月%d日", time.localtime(time.mktime(time.localtime()) - (3600 * 24 * 2)))
        elif art_time.find("天前") != -1:
            add_date = int(art_time.split("天前")[0])
            art_r_time = time.strftime("%m月%d日", time.localtime(time.mktime(time.localtime()) - (3600 * 24 * add_date)))
        elif art_time.find("小时前") != -1 or art_time.find("分钟前") != -1:
            art_r_time = time.strftime("%m月%d日", time.localtime())
        else:
            art_r_time = art_time
        return art_r_time

    def save_article(self, parent_dir, art_r_time, art_y_time, art_content=None):
        """
        保存文章
        :param parent_dir:父目录
        :param art_r_time:日期
        :param art_y_time:年份
        :param art_content:文章内容
        :return:返回标题作为图片名
        """
        document = Document()
        if not os.path.exists(parent_dir + "/" + art_y_time):
            os.makedirs(parent_dir + "/" + art_y_time)
        if not os.path.exists(parent_dir + "/" + art_y_time + "/" + art_r_time):
            os.makedirs(parent_dir + "/" + art_y_time + "/" + art_r_time)
        if art_content is None:
            docxname = art_r_time + str(random.randint(0, 9999)) + "_"
        else:
            art_content = art_content
            art_list = art_content.split("\n")
            if len(art_list) > 1:
                if len(art_list[0]) > 20:
                    docxname = art_list[0].split("，")[0]
                    self.w_docx_content(document, art_list[1:])
                else:
                    docxname = art_list[0]
                    self.w_docx_title(document, art_list[0])
                    self.w_docx_content(document, art_list[1:])
            else:
                if len(art_content) > 20:
                    docxname = art_content.split("，")[0]
                else:
                    docxname = art_content
                self.w_docx_content(document, art_content.split("\n"))
            docxname = docxname.replace("\"", "”").replace("\"", "“").replace(".", "。").replace(",", "，")
            document.save(parent_dir + "/" + art_y_time + '/' + art_r_time + '/' + docxname + '.docx')
        return docxname + "_", parent_dir + "/" + art_y_time + '/' + art_r_time + '/'

    def download_media(self):
        """
        下载媒体资源
        :return: 资源存放路径和资源类型
        """
        try:
            while True:
                try:
                    time_tag = self.driver.find_element(By.ID, 'com.tencent.mm:id/ng')
                    break
                except NoSuchElementException:
                    self.driver.swipe(FLICK_START_X, FLICK_START_Y + 500, FLICK_START_X, FLICK_START_Y)
            self.driver.find_element(By.ID, 'com.tencent.mm:id/ms').click()
            sleep(0.2)
            try:
                # 这条语句用来触发异常
                self.driver.find_element(By.ID, 'com.tencent.mm:id/gvo')
                items = self.wait.until(EC.presence_of_all_elements_located((By.ID, "com.tencent.mm:id/gvo")))
                for i in range(len(items)):
                    self.driver.swipe(FLICK_START_X, FLICK_START_Y, FLICK_START_X, FLICK_START_Y, 2000)
                    sleep(0.2)
                    self.driver.find_element(By.XPATH, '//*[@text="保存图片"]').click()
#                   self.driver.find_element(By., 'com.tencent.mm:id/f15').click()
                    sleep(0.2)
                    self.driver.swipe(FLICK_START_X + 600, FLICK_START_Y, FLICK_START_X, FLICK_START_Y)
                    sleep(0.2)
                else:
                    self.driver.tap([(FLICK_START_X, FLICK_START_Y)])
            except NoSuchElementException:
                self.driver.swipe(FLICK_START_X, FLICK_START_Y, FLICK_START_X, FLICK_START_Y, 2000)
                sleep(0.2)
                self.driver.find_element(By.XPATH, '//*[@text="保存图片"]').click()
                sleep(0.2)
                self.driver.tap([(FLICK_START_X, FLICK_START_Y)])
                pass
            return RP_DIRECTORY, '.jpg'
        except NoSuchElementException:
            self.driver.find_element(By.ID, 'com.tencent.mm:id/b47').click()
            sleep(0.5)
            self.driver.swipe(FLICK_START_X, FLICK_START_Y, FLICK_START_X, FLICK_START_Y, 2000)
            sleep(0.2)
            self.driver.find_element(By.XPATH, '//*[@text="保存视频"]').click()
            sleep(0.2)
            self.driver.tap([(FLICK_START_X, FLICK_START_Y)])
            return RV_DIRECTORY, '.mp4'
            pass

    def save_media(self, l_file, r_directory, l_directory, media_type):
        """
        导出媒体资源
        :param r_directory: 远程文件路径
        :param l_directory: 本地保存文件路径
        :param media_type: 媒体资源文件类型
        :param l_file: 保存文件名
        :return:
        """
        media_list = self.adb_device.shell('ls ' + r_directory + ' |cat').split("\n")[:-1]
        i = 1
        for r_file in media_list:
            self.adb_device.pull(r_directory + r_file, l_directory + l_file + '_' + str(i) + media_type)
            self.adb_device.shell('rm -f ' + r_directory + r_file)
            i += 1

    def craw(self):
        """
        抓取
        :return:
        """

        #        df = pd.read_excel("venv/data/wechat.xlsx")

        # 滑动到结尾
        end_point1 = "com.tencent.mm:id/g39"
        end_point2 = "com.tencent.mm:id/ifi"
        end_point3 = False
        # 年份元素
        years = "com.tencent.mm:id/jxl"
        # 时间元素
        wdate = "com.tencent.mm:id/ju9"
        # 月份元素
        mdate = "com.tencent.mm:id/juc"
        # 日期元素
        ddate = "com.tencent.mm:id/jsu"
        # 含图文字内容
        psms = "com.tencent.mm:id/c22"
        # 含链接文字内容和纯文字内容
        wsms = "com.tencent.mm:id/c2h"
        # 链接标题
        llink = "com.tencent.mm:id/kpq"
        # 包含图片
        has_picture = "com.tencent.mm:id/ju8"
        ryear_list = []
        rdate_list = []
        rtext_list = []
        rlink_list = []
        dict1 = {}
        tyear = str(time.localtime().tm_year) + "年"
        rdate = ' '
        if len(TDATETIME) == 0:
            ttime = time.mktime(time.strptime("1900年1月1日", "%Y年%m月%d日"))
        else:
            ttime = time.mktime(time.strptime(TDATETIME, "%Y年%m月%d日"))
        sleep(SCROLL_SLEEP_TIME)
        while True:
            # 当前页面显示的所有目标内容集合
            #            items = self.wait.until(EC.presence_of_all_elements_located((By.XPATH, '//*[@resource-id="com.tencent.mm:id/br8"]')))
            items = self.wait.until(EC.presence_of_all_elements_located((By.ID, "com.tencent.mm:id/br8")))
            # 遍历每条内容
            for item in items:
                rtext = ' '
                ryear = ' '
                rlink = ' '
                try:
                    ryear = item.find_element(By.ID, years).get_attribute('text')
                    dict2 = {}
                    dict2['年份'] = ryear_list
                    dict2['日期'] = rdate_list
                    dict2['内容'] = rtext_list
                    dict2['链接'] = rlink_list
                    df2 = pd.DataFrame(dict2)
                    with pd.ExcelWriter('F:/test.xlsx', mode='a', engine='openpyxl') as writer:
                        writer.if_sheet_exists = "replace"
                        df1 = pd.read_excel('F:/test.xlsx', sheet_name='Sheet1', index_col=0)
                        f = [df1, df2]
                        result = pd.concat(f, axis=0)
                        result.to_excel(writer, sheet_name="Sheet1", index_label=0)
                    ryear_list = []
                    rdate_list = []
                    rtext_list = []
                    rlink_list = []
                    tyear = ryear
                #    print(ryear + "+++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++++")
                except NoSuchElementException:
                    pass
                try:
                    rtext = item.find_element(By.ID, psms).get_attribute('text')
                    try:
                        a = rtext_list.index(rtext)
                        continue
                    except ValueError:
                        pass
                except NoSuchElementException:
                    pass
                try:
                    rtext = item.find_element(By.ID, wsms).get_attribute('text')
                    try:
                        a = rtext_list.index(rtext)
                        continue
                    except ValueError:
                        pass
                except NoSuchElementException:
                    pass
                try:
                    rlink = item.find_element(By.ID, llink).get_attribute('text')
                    try:
                        a = rtext_list.index(rlink)
                        continue
                    except ValueError:
                        pass
                except NoSuchElementException:
                    pass
                try:
                    rdate = item.find_element(By.ID, ddate).get_attribute('text')
                    try:
                        rdate1 = item.find_element(By.ID, mdate).get_attribute('text')
                        rdate = rdate1 + rdate + "日"
                        ntime = time.mktime(time.strptime(tyear + rdate, "%Y年%m月%d日"))
                        print(ntime)
                        if ntime < ttime:
                            end_point3 = True
                            break
                    except NoSuchElementException:
                        print("没取到月份")
                        pass
                except NoSuchElementException:
                    pass
                rdate_list.append(rdate)
                rlink_list.append(rlink)
                rtext_list.append(rtext)
                ryear_list.append(ryear)
            bounds = items[len(items) - 1].get_attribute('bounds')
            m = re.findall(r'\d+', bounds)
            print(m)
            print(type(m))
            half_m = int(m[len(m) - 1]) / 2 - 100
            # 判断是否已经到结尾
            page = self.driver.page_source
            if end_point1 in page or end_point2 in page or end_point3:
                break
            # 上滑
            self.driver.swipe(FLICK_START_X, FLICK_START_Y + half_m, FLICK_START_X, FLICK_START_Y, 2000)
            self.driver.swipe(FLICK_START_X, FLICK_START_Y + half_m, FLICK_START_X, FLICK_START_Y, 2000)
        dict1['年份'] = ryear_list
        dict1['日期'] = rdate_list
        dict1['内容'] = rtext_list
        dict1['链接'] = rlink_list
        print(dict1)
        df = pd.DataFrame(dict1)
        #    df.to_excel('venv/data/wechat.xlsx', sheet_name='sheet1')
        with pd.ExcelWriter('F:/test.xlsx', mode='a', engine='openpyxl') as writer:
            writer.if_sheet_exists = "replace"
            df1 = pd.read_excel('F:/test.xlsx', sheet_name='Sheet1', index_col=0)
            f = [df1, df]
            result = pd.concat(f, axis=0)
            result.to_excel(writer, sheet_name="Sheet1", index_label=0)

    def craw_all(self):
        """
        抓取
        :return:
        """

        # 滑动到结尾
        end_point1 = "com.tencent.mm:id/g39"
        end_point2 = "com.tencent.mm:id/ifi"
        end_point3 = False
        # 年份元素
        e_years = "com.tencent.mm:id/jxl"
        # 时间元素
        e_wdate = "com.tencent.mm:id/ju9"
        # 月份元素
        e_mdate = "com.tencent.mm:id/juc"
        # 日期元素
        e_ddate = "com.tencent.mm:id/jsu"
        # 含图文字内容
        e_psms = "com.tencent.mm:id/c22"
        # 含链接文字内容和纯文字内容
        e_wsms = "com.tencent.mm:id/c2h"
        # 链接标题
        e_llink = "com.tencent.mm:id/kpq"
        # 包含图片
        e_has_picture = "com.tencent.mm:id/ju8"

#        tyear = str(time.localtime().tm_year) + "年"
        tyear = self.now_year
        now_time = time.mktime(time.localtime())
        rdate = ' '
        if len(TDATETIME) == 0:
            ttime = time.mktime(time.strptime("1900年1月1日", "%Y年%m月%d日"))
        else:
            ttime = time.mktime(time.strptime(TDATETIME, "%Y年%m月%d日"))
        sleep(SCROLL_SLEEP_TIME)
        pyq_list = deque()
        while True:
            # 当前页面显示的所有目标内容集合
            #            items = self.wait.until(EC.presence_of_all_elements_located((By.XPATH, '//*[@resource-id="com.tencent.mm:id/br8"]')))
            items = self.wait.until(EC.presence_of_all_elements_located((By.ID, "com.tencent.mm:id/br8")))
            # 遍历每条内容
            for item in items:
                if item.id in pyq_list:
                    continue
                elif item.location['y'] < 328:
                    continue
                elif item.location['y'] + item.size['height'] > 1840:
                    continue
                else:
                    pyq_list.append(item.id)
                    if len(pyq_list) > 6:
                        pyq_list.popleft()
                rtext = None
                ryear = None
                rlink = ''
                try:
                    item.find_element(By.XPATH, '//*[@content-desc="拍照分享"]')
                    continue
                except NoSuchElementException:
                    pass
                try:
                    ryear = item.find_element(By.ID, e_years).get_attribute('text')
                    tyear = ryear
                except NoSuchElementException:
                    ryear = tyear
                    pass
                try:
                    rtext = item.find_element(By.ID, e_psms).get_attribute('text')
                except NoSuchElementException:
                    pass
                try:
                    rtext = item.find_element(By.ID, e_wsms).get_attribute('text')
                except NoSuchElementException:
                    pass
                """
                try:
                    rlink = item.find_element(By.ID, e_llink).get_attribute('text')
                    item.find_element(By.ID, e_llink).click()
                    sleep(0.5)
                    self.driver.find_element(By.ID, e_llink).click()
                    sleep(5)
#                    page = self.wait.until(EC.presence_of_element_located((By.XPATH, '//android.widget.ImageView[@content-desc="更多信息"]')))
#                    page.click()
                    ct = self.driver.contexts
                    print(ct)
                    self.driver.switch_to.context('WEBVIEW_com.tencent.mm:toolsmp')
                    self.driver.find_element(By.XPATH, '//android.widget.ImageView[@content-desc="更多信息"]').click()
                    sleep(0.5)
                    self.driver.find_element(By.XPATH, '//*[@text="复制链接"]').click()
                    print(self.driver.get_clipboard_text())
                    self.driver.switch_to.default_content()
                except NoSuchElementException:
                    pass
                """
                try:
                    rdate = item.find_element(By.ID, e_ddate).get_attribute('text')
                    try:
                        rdate1 = item.find_element(By.ID, e_mdate).get_attribute('text')
                        rdate = rdate1 + rdate + "日"
                    except NoSuchElementException:
                        rdate = self.tranform_time(rdate)
                        pass
                    ntime = time.mktime(time.strptime(ryear + rdate, "%Y年%m月%d日"))
                    now_time = ntime
                    if ntime < ttime:
                        end_point3 = True
                        break
                except NoSuchElementException:
                    pass
                if now_time > self.start_time:
                    continue
                mtext, l_directory = self.save_article(L_DIRECTORY, rdate, ryear, rtext)
                # 下载媒体资源
                try:
                    item.find_element(By.ID, e_has_picture).click()
                    sleep(0.5)
                    r_directory, media_type = self.download_media()
                    self.save_media(mtext, r_directory, l_directory, media_type)
                    sleep(0.5)
                    self.driver.find_element(By.XPATH, '//*[@content-desc="返回"]').click()
                    sleep(0.5)
                except NoSuchElementException:
                    pass
            bounds = items[len(items) - 1].get_attribute('bounds')
            m = re.findall(r'\d+', bounds)
            half_m = int(m[len(m) - 1]) / 2 - 100
            # 判断是否已经到结尾
            page = self.driver.page_source
            if end_point1 in page or end_point2 in page or end_point3:
                break
            # 上滑
            self.driver.swipe(FLICK_START_X, FLICK_START_Y + half_m, FLICK_START_X, FLICK_START_Y, 2000)
            # self.driver.swipe(FLICK_START_X, FLICK_START_Y + half_m, FLICK_START_X, FLICK_START_Y, 2000)
            # self.driver.swipe(FLICK_START_X, FLICK_START_Y + half_m, FLICK_START_X, FLICK_START_Y, 2000)


    def main(self):
        """
        入口
        :return:
        """
        # 通讯录
        self.contact()
        # 爬取
        self.craw_all()


if __name__ == '__main__':
    moments = Moments()
    moments.main()
