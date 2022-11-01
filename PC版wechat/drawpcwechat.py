import re
import time
import psutil
import pywinauto
from pywinauto.application import Application
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import random
from collections import deque
import json


class DrawPCWechat:
    def __init__(self):
        """
        初始化
        """
        self.PID = 0
        for proc in psutil.process_iter():
            try:
                pinfo = proc.as_dict(attrs=['pid', 'name'])
            except psutil.NoSuchProcess:
                pass
            else:
                if 'WeChat.exe' == pinfo['name']:
                    self.PID = pinfo['pid']
        self.app = Application(backend='uia').connect(process=self.PID)
        self.app_win32 = Application().connect(process=self.PID)
        with open("./config.json", "r", encoding="utf-8") as config_data:
            config_dict = json.load(config_data)
        self.username = config_dict['username']
        self.ac_time = config_dict['action_time'].split("年")[1]
        self.ac_year = config_dict['action_time'][:5]
        self.now_year = time.strftime("%Y年", time.localtime())
        self.parent_path = config_dict['parent_path']

    def click_btn(self, winsf):
        """
        点击按钮
        :param winsf:
        :return:
        """
        btn_cords = winsf.rectangle()
        pywinauto.mouse.click(button='left', coords=(int((btn_cords.left + btn_cords.right) / 2), int((btn_cords.top + btn_cords.bottom) / 2)))

    def w_docx_content(self, document, art_list):
        """
        写入段落
        :param document:
        :param art_list:
        :return:
        """
        for p in art_list:
            pg = document.add_paragraph()
            # 设置内容
            pg.text = p.replace(" ", "")
            # 设置字号
            pg.style.font.size = Pt(14)
            pg.paragraph_format.first_line_indent = pg.style.font.size * 2
            # 设置英文字符字体
            pg.style.font.name = '黑体'
            # 设置中文字符字体
            pg.style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    def w_docx_title(self, document, art_title):
        """
        写入标题
        :param document:
        :param art_title:
        :return:
        """
        run = document.add_heading("", level=0).add_run(art_title.lstrip())
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        run.font.size = Pt(14)

    def tranform_time(self, art_time):
        """
        修改时间格式
        :param art_time:
        :return:
        """
        art_y_time = self.now_year
        if art_time == "昨天":
            art_r_time = time.strftime("%m月%d日", time.localtime(time.mktime(time.localtime()) - (3600 * 24)))
        elif art_time == "前天":
            art_r_time = time.strftime("%m月%d日", time.localtime(time.mktime(time.localtime()) - (3600 * 24 * 2)))
        elif art_time.find("天前") != -1:
            add_date = int(art_time.split("天前")[0])
            art_r_time = time.strftime("%m月%d日", time.localtime(time.mktime(time.localtime()) - (3600 * 24 * add_date)))
        elif art_time.find("小时前") != -1 or art_time.find("分钟前") != -1:
            art_r_time = time.strftime("%m月%d日", time.localtime())
        elif art_time.find("年") != -1:
            art_y_time = art_time[:5]
            art_r_time = art_time.split(" ")[0].split("年")[1]
        else:
            art_r_time = art_time.split(" ")[0]
        return art_r_time, art_y_time

    def save_article(self, parent_dir, art_time, u_name, art_content=None):
        """
        保存文章
        :param parent_dir:
        :param art_time:
        :param art_content:
        :return:返回标题作为图片名
        """
        document = Document()
        art_r_time = art_time
        art_r_time, art_y_time = self.tranform_time(art_time)
        if not os.path.exists(parent_dir + "/" + u_name + "/" + art_y_time):
            os.makedirs(parent_dir + "/" + u_name + "/" + art_y_time)
        if not os.path.exists(parent_dir + "/" + u_name + "/" + art_y_time + "/" + art_r_time):
            os.makedirs(parent_dir + "/" + u_name + "/" + art_y_time + "/" + art_r_time)
        if art_content is None:
            docxname = art_r_time + str(random.randint(0, 9999)) + "_"
        else:
            art_list = art_content.split("\n")
            if len(art_list) > 1:
                if len(art_list[0]) > 20:
                    docxname = art_list[0].split("，")[0]
                    self.w_docx_content(document, art_list[1:])
                else:
                    docxname = art_list[0]
                    self.w_docx_title(document, art_list[0])
                    self.w_docx_content(document, art_list[1:])
            else:
                if len(art_content) > 20:
                    docxname = art_content.split("，")[0]
                else:
                    docxname = art_content
                self.w_docx_content(document, art_content.split("\n"))
            document.save(parent_dir + "/" + u_name + "/" + art_y_time + '/' + art_r_time + '/' + docxname + '.docx')
        return docxname + "_", parent_dir + "/" + u_name + "/" + art_y_time + '/' + art_r_time

    def open_pyq(self):
        """
        打开朋友圈
        :return:
        """
        pyq_win = self.app['朋友圈']
        if not pyq_win.exists():
            win = self.app['微信']
            win.set_focus()
            pyq_btn = win.child_window(title="朋友圈", control_type="Button")
            cords = pyq_btn.rectangle()
            pywinauto.mouse.click(button='left', coords=(cords.left + 10, cords.top + 10))
            pyq_win = self.app['朋友圈']
        pyq_win.set_focus()
        pyq_cords = pyq_win.rectangle()
        print(pyq_cords)
        pywinauto.mouse.press(button='left', coords=(int((pyq_cords.left + pyq_cords.right)/2), pyq_cords.bottom - 2))
        # pywinauto.mouse.move(coords=(int((pyq_cords.left + pyq_cords.right)/2), pyq_cords.bottom + 300))
        time.sleep(2)
        pywinauto.mouse.release(button='left', coords=(int((pyq_cords.left + pyq_cords.right)/2), pyq_cords.bottom + 398))
        print("打开朋友圈...")

    def locate_pyq(self):
        time.sleep(1)
        pyq_win = self.app['朋友圈']
        pyq_cords = pyq_win.rectangle()
        p_time = pyq_win.child_window(title_re="\d+分钟前|\d+小时前|昨天|前天|\d天前|.*月\d+日.\d+:\d+", control_type="Text", found_index=0).texts()[0]
        p_r_time, p_y_time = self.tranform_time(p_time)
        if time.mktime(time.strptime(p_y_time + p_r_time, "%Y年%m月%d日")) > time.mktime(
                time.strptime(self.ac_year + self.ac_time, "%Y年%m月%d日")):
            pyq_win.set_focus()
            pywinauto.mouse.scroll(
                coords=(int((pyq_cords.left + pyq_cords.right) / 2), int((pyq_cords.top + pyq_cords.bottom) / 2)),
                wheel_dist=-5)
            return False
        else:
            return True



    def draw_pyq(self):
        """
        爬取朋友圈
        :return:
        """
        pyq_win = self.app['朋友圈']
        pyq_cords = pyq_win.rectangle()
        pyq_list = deque()
        scroll_cords = 0
        print("开始抓取数据...")
        while True:
            i_list = pyq_win.child_window(title="朋友圈", control_type="List").children(control_type="ListItem")
            i_list_count = len(pyq_win.child_window(title="朋友圈", control_type="List").children(control_type="ListItem"))
            # pyq_win.child_window(title="朋友圈", control_type="List").dump_tree()
            j = 0
            jr = 0
            while j < i_list_count:
                w_list = pyq_win.child_window(control_type="ListItem", found_index=jr)
                try:
                    w_list_cords = w_list.rectangle()
                except pywinauto.findwindows.ElementNotFoundError:
                    break
                if w_list_cords.top < pyq_cords.top or w_list_cords.bottom > pyq_cords.bottom:
                    # 对象超过窗口范围
                    j += 1
                    jr += 1
                    continue
                if w_list.parent() != pyq_win.child_window(title="朋友圈", control_type="List").wrapper_object():
                    # scroll_cords = w_list_cords.bottom - pyq_cords.top
                    # 非子对象
                    jr += 1
                    continue
#                u_name = str(w_list.wrapper_object()).split("\n")[0].split("\'")[1]
                if str(w_list.wrapper_object()) in pyq_list:
                    scroll_cords = w_list_cords.bottom - pyq_cords.top
                    j += 1
                    jr += 1
                    # 已存在元素，跳过
                    continue
                else:
                    # 添加元素进队列
                    pyq_list.append(str(w_list.wrapper_object()))
                    if len(pyq_list) > 5:
                        pyq_list.popleft()
                u_name = str(w_list.wrapper_object()).split("\n")[0].split("\'")[1]
                if u_name != self.username:
                    j += 1
                    jr += 1
                    continue
                try:
                    f_time = w_list.child_window(control_type="Text", found_index=1).texts()[0]
                    if f_time.find("视频号") != -1:
                        j += 1
                        jr += 1
                        continue
                    f_text = w_list.child_window(control_type="Text", found_index=0).texts()[0]
                    docx_name, save_path = self.save_article(self.parent_path, f_time, u_name, f_text)
                except pywinauto.findwindows.ElementNotFoundError:
                    f_time = w_list.child_window(control_type="Text", found_index=0).texts()[0]
                    docx_name, save_path = self.save_article(self.parent_path, f_time, u_name)

                # 处理图片
                if w_list.child_window(title_re="包含\d+张图片", control_type="Pane").exists():
                    p_num = w_list.child_window(title_re="包含\d+张图片", control_type="Pane").texts()[0]
                    pp_num = int(re.findall("\d+", p_num)[0])
                    # w_list.child_window(title_re="包含\d+张图片", control_type="Pane").dump_tree()
                    p = 0
                    while p < pp_num:
                        self.click_btn(w_list.child_window(title_re="包含\d+张图片", control_type="Pane").child_window(title="图片", control_type="Pane", found_index=p))
                        tupian_win = self.app['图片查看']
                        time.sleep(0.5)
                        # 如果图片宽度过小，另存为按钮会被隐藏
                        try:
                            self.click_btn(tupian_win.child_window(title="另存为...", control_type="Button"))
                        except pywinauto.findwindows.ElementNotFoundError:
                            self.click_btn(tupian_win.child_window(title="更多", control_type="Button"))
                            self.click_btn(tupian_win.child_window(title="另存为...", control_type="MenuItem"))
                        lingcunwei_win = self.app_win32['另存为...']
                        lingcunwei_win.child_window(title_re="微信图片_.*", class_name="ComboBox").child_window(
                            title_re="微信图片_.*", class_name="Edit").set_text(docx_name + str(p + 1) + ".jpg")
                        lcw_addr_list = \
                        lingcunwei_win.child_window(class_name="msctls_progress32").child_window(title_re="地址: .*", class_name="ToolbarWindow32").texts()[
                            0]
                        lcw_addr = lcw_addr_list.split("地址: ")[1]
                        lingcunwei_win.child_window(class_name="msctls_progress32").child_window(title_re="地址: .*", class_name="ToolbarWindow32").click()
                        lingcunwei_win.child_window(class_name="msctls_progress32").child_window(title=lcw_addr, class_name="Edit").set_text(save_path)
                        time.sleep(0.5)
                        lingcunwei_win.type_keys("{ENTER}")
                        time.sleep(0.5)
                        lingcunwei_win.type_keys("%{s}")
#                        time.sleep(0.5)
#                        try:
#                            self.click_btn(self.app_win32['确认另存为'].child_window(title="否(&N)", class_name="Button"))
#                        except pywinauto.findwindows.ElementNotFoundError:
#                            pass
                        time.sleep(0.5)
                        self.click_btn(tupian_win.child_window(title="关闭", control_type="Button", found_index=0))
                        p += 1

                # 处理视频
                if w_list.child_window(title="视频", control_type="Pane").exists():
                    self.click_btn(w_list.child_window(title="视频", control_type="Pane"))
                    shipin_win = self.app['视频查看']
                    if not shipin_win.exists():
                        self.click_btn(w_list.child_window(title="视频", control_type="Pane"))
                    video_cords = shipin_win.rectangle()
                    pywinauto.mouse.click(button='right', coords=(int((video_cords.left + video_cords.right) / 2), int((video_cords.top + video_cords.bottom) / 2)))
                    time.sleep(0.5)
                    self.click_btn(shipin_win.child_window(title="另存为...", control_type="MenuItem"))
                    lingcunwei_win = self.app_win32['另存为...']
                    lingcunwei_win.child_window(title_re=".*\.mp4", class_name="ComboBox", found_index=0).child_window(
                        title_re=".*\.mp4", class_name="Edit").set_text(docx_name + ".mp4")
                    lcw_addr_list = \
                    lingcunwei_win.child_window(class_name="msctls_progress32").child_window(title_re="地址: .*", class_name="ToolbarWindow32").texts()[0]
                    lcw_addr = lcw_addr_list.split("地址: ")[1]
                    lingcunwei_win.child_window(class_name="msctls_progress32").child_window(title_re="地址: .*", class_name="ToolbarWindow32").click()
                    lingcunwei_win.child_window(class_name="msctls_progress32").child_window(title=lcw_addr, class_name="Edit").set_text(save_path)
                    lingcunwei_win.type_keys("{ENTER}")
                    time.sleep(1)
                    lingcunwei_win.type_keys("%{s}")
                    self.click_btn(shipin_win.child_window(title="关闭", control_type="Button", found_index=0))
                scroll_cords = w_list_cords.bottom - pyq_cords.top
                j += 1
                jr += 1

            pyq_win.set_focus()
            pywinauto.mouse.scroll(coords=(int((pyq_cords.left + pyq_cords.right) / 2), int((pyq_cords.top + pyq_cords.bottom) / 2)), wheel_dist=-int(scroll_cords/150))
            time.sleep(1)

    def main(self):
        """
        入口
        :return:
        """
        input("按回车键开始任务...")
        self.open_pyq()
        while True:
            if self.locate_pyq():
                print("完成定位...")
                self.draw_pyq()


if __name__ == '__main__':
    drawpcwechat = DrawPCWechat()
    drawpcwechat.main()
